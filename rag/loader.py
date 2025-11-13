# -*- coding: utf-8 -*-
# @File: loader.py
# @Author: yaccii
# @Time: 2025-11-09 23:10
# @Description:
import io
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

from fastapi import UploadFile


@dataclass
class LoaderConfig:
    max_file_size_mb: int = 25

    csv_preview_rows: int = 200
    csv_preview_cols: int = 50

    excel_preview_rows: int = 200
    excel_preview_cols: int = 30

    html_extractor: str = "trafilatura"  # trafilatura | bs4
    text_max_chars: int = 2_000_000


DEFAULT_CONFIG = LoaderConfig()
_SENT_END = r"[。！？!?…：:；;．.]"  # 视需要再加


# -------------------- utils --------------------
def _unwrap_soft_linebreaks(text: str) -> str:
    """
    合并段内软换行：
    - 单换行且前面不是句末标点，则合并为空格；
    - 保留空行（段落分隔）
    - 处理被硬拆开的函数名/下划线：把换行去掉
    """
    # 1) 把形如 "get_num_\n\ntokens" 里的换行去掉（函数/下划线拼写）
    text = re.sub(r"(_)\n+([A-Za-z0-9])", r"\1\2", text)

    # 2) 段内软换行：前后都是非空且前一个字符不是句末标点 -> 用空格合并
    def _merge(m: re.Match) -> str:
        prev = m.group(1)
        nextc = m.group(2)
        if re.search(_SENT_END + r"$", prev):
            return prev + "\n" + nextc  # 真换行
        return prev + " " + nextc  # 软换行 => 空格

    # 用占位捕获前后字符，避免误删段落空行
    text = re.sub(r"([^\n])\n([^\n])", _merge, text)

    # 3) 多个空行压缩为一个空行（保留段落感）
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _normalize_newlines(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _ensure_limit(text: str, limit: int, suffix: str = "\n\n[... truncated ...]") -> str:
    cleaned = _unwrap_soft_linebreaks(_normalize_newlines(text or ""))

    if not isinstance(limit, int) or limit <= 0:
        return ""

    if len(cleaned) > limit:
        return cleaned[:limit].rstrip() + suffix

    return cleaned


def _read_text_guess_encoding(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="strict")
    except:
        try:
            import chardet
        except Exception as e:
            raise RuntimeError("Please install 'chardet' to guess text encoding") from e
        raw = path.read_bytes()
        det = chardet.detect(raw) or {}
        enc = det.get("encoding") or "utf-8"
        return raw.decode(enc, errors="ignore")


def _pformat_json(obj) -> str:
    return json.dumps(obj, indent=2, ensure_ascii=False)


# -------------------- concrete loaders --------------------

def _load_txt(path: Path) -> str:
    return _read_text_guess_encoding(path)


def _load_md(path: Path) -> str:
    return _read_text_guess_encoding(path)


def _load_csv_like(path: Path, cfg: LoaderConfig) -> str:
    try:
        import pandas as pd
    except Exception as e:
        raise RuntimeError("Please install 'pandas' to parse CSV/TSV") from e

    sep = "," if path.suffix.lower() == ".csv" else "\t"
    df = pd.read_csv(
        str(path), sep=sep, nrows=cfg.csv_preview_rows, encoding="utf-8", engine="python", on_bad_lines="skip"
    )
    if df.shape[1] > cfg.csv_preview_cols:
        df = df.iloc[:, : cfg.csv_preview_cols]
    buf = io.StringIO()
    df.to_csv(buf, index=False, sep="\t")
    return buf.getvalue()


def _load_json(path: Path) -> str:
    obj = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
    return _pformat_json(obj)


def _load_pdf(path: Path) -> str:
    try:
        import fitz  # pymupdf
    except Exception as e:
        raise RuntimeError("Please install 'pymupdf' to parse PDF") from e
    texts: List[str] = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            texts.append(page.get_text("text"))
    return "\n".join(texts)


def _load_docx(path: Path) -> str:
    try:
        import docx
    except Exception as e:
        raise RuntimeError("Please install 'python-docx' to parse .docx") from e
    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs)


def _load_excel(path: Path, cfg: LoaderConfig) -> str:
    try:
        import pandas as pd
    except Exception as e:
        raise RuntimeError("Please install 'pandas' to parse Excel") from e

    xls = pd.ExcelFile(str(path))
    out: List[str] = []
    for name in xls.sheet_names[:10]:
        df = pd.read_excel(xls, sheet_name=name, nrows=cfg.excel_preview_rows, engine=None)
        if df.shape[1] > cfg.excel_preview_cols:
            df = df.iloc[:, : cfg.excel_preview_cols]
        buf = io.StringIO()
        df.to_csv(buf, index=False, sep="\t")
        out.append(f"=== Sheet: {name} ===\n{buf.getvalue()}")
    return "\n\n".join(out)


def _strip_html_basic(html: str) -> str:
    try:
        from bs4 import BeautifulSoup
    except Exception as e:
        raise RuntimeError("Please install 'beautifulsoup4' to parse HTML") from e
    soup = BeautifulSoup(html, "lxml")  # 需要 lxml
    for tag in soup(["script", "style"]):
        tag.decompose()
    return "\n".join(s.strip() for s in soup.stripped_strings)


def _load_html_from_string(html: str, *, prefer: str = "trafilatura") -> Tuple[Optional[str], str]:
    title = None
    content = ""
    if prefer.lower() == "trafilatura":
        try:
            import trafilatura
            from lxml import html as lxml_html
            doc = lxml_html.fromstring(html)
            tnodes = doc.xpath("//title")
            if tnodes and tnodes[0].text:
                title = (tnodes[0].text or "").strip()
            extracted = trafilatura.extract(
                html, include_comments=False, include_tables=True, favor_recall=True
            )
            content = extracted or _strip_html_basic(html)
        except Exception:
            content = _strip_html_basic(html)
    else:
        content = _strip_html_basic(html)
    return title, content


# -------------------- public helpers --------------------

def load_text_from_file_path(file_path: str, *, config: Optional[LoaderConfig] = None) -> Tuple[str, str]:
    cfg = config or DEFAULT_CONFIG
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(str(path))

    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > cfg.max_file_size_mb:
        raise ValueError(f"File too large: {size_mb:.1f}MB > {cfg.max_file_size_mb}MB")

    suffix = path.suffix.lower()
    title = path.stem

    if suffix in {".txt", ".log"}:
        content = _load_txt(path)
        return title, _ensure_limit(_normalize_newlines(content), cfg.text_max_chars)

    if suffix in {".md", ".rst"}:
        content = _load_md(path)
        return title, _ensure_limit(_normalize_newlines(content), cfg.text_max_chars)

    if suffix in {".csv", ".tsv"}:
        content = _load_csv_like(path, cfg)
        return title, _ensure_limit(_normalize_newlines(content), cfg.text_max_chars)

    if suffix == ".json":
        content = _load_json(path)
        return title, _ensure_limit(_normalize_newlines(content), cfg.text_max_chars)

    if suffix == ".pdf":
        content = _load_pdf(path)
        return title, _ensure_limit(_normalize_newlines(content), cfg.text_max_chars)

    if suffix == ".docx":
        content = _load_docx(path)
        return title, _ensure_limit(_normalize_newlines(content), cfg.text_max_chars)

    if suffix in {".xlsx", ".xls"}:
        content = _load_excel(path, cfg)
        return title, _ensure_limit(_normalize_newlines(content), cfg.text_max_chars)

    if suffix in {".html", ".htm"}:
        html = path.read_text(encoding="utf-8", errors="ignore")
        t, content = _load_html_from_string(html, prefer=cfg.html_extractor)
        return (t or title), _ensure_limit(_normalize_newlines(content), cfg.text_max_chars)

    raise ValueError(f"Unsupported file type: {suffix}")


async def load_text_from_upload_file(file: UploadFile, *, config: Optional[LoaderConfig] = None) -> Tuple[str, str]:
    cfg = config or DEFAULT_CONFIG
    filename = file.filename or "upload"
    suffix = Path(filename).suffix.lower()
    title = Path(filename).stem

    # 尽量避免多次读取；这里将其读入内存（小文件友好），大文件建议你在路由层先落盘再走 file_path 版本
    raw = await file.read()

    if len(raw) > cfg.max_file_size_mb * 1024 * 1024:
        raise ValueError(f"File too large: {(len(raw) / 1024 / 1024):.1f}MB > {cfg.max_file_size_mb}MB")

    if suffix in {".txt", ".log", ".md", ".rst"}:
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            try:
                import chardet
                enc = (chardet.detect(raw) or {}).get("encoding") or "utf-8"
                text = raw.decode(enc, errors="ignore")
            except Exception as e:
                raise RuntimeError("Please install 'chardet' to guess text encoding") from e
        return title, _ensure_limit(_normalize_newlines(text), cfg.text_max_chars)

    if suffix in {".csv", ".tsv"}:
        try:
            import pandas as pd
        except Exception as e:
            raise RuntimeError("Please install 'pandas' to parse CSV/TSV") from e
        sep = "," if suffix == ".csv" else "\t"
        df = pd.read_csv(io.BytesIO(raw), sep=sep, nrows=cfg.csv_preview_rows, encoding="utf-8", engine="python",
                         on_bad_lines="skip")
        if df.shape[1] > cfg.csv_preview_cols:
            df = df.iloc[:, : cfg.csv_preview_cols]
        buf = io.StringIO()
        df.to_csv(buf, index=False, sep="\t")
        return title, _ensure_limit(_normalize_newlines(buf.getvalue()), cfg.text_max_chars)

    if suffix == ".json":
        try:
            obj = json.loads(raw.decode("utf-8", errors="ignore"))
        except Exception:
            obj = {}
        return title, _ensure_limit(_normalize_newlines(_pformat_json(obj)), cfg.text_max_chars)

    if suffix == ".pdf":
        try:
            import fitz
        except Exception as e:
            raise RuntimeError("Please install 'pymupdf' to parse PDF") from e
        texts: List[str] = []
        with fitz.open(stream=raw, filetype="pdf") as doc:
            for page in doc:
                texts.append(page.get_text("text"))
        return title, _ensure_limit(_normalize_newlines("\n".join(texts)), cfg.text_max_chars)

    if suffix == ".docx":
        try:
            import docx
        except Exception as e:
            raise RuntimeError("Please install 'python-docx' to parse .docx") from e
        bio = io.BytesIO(raw)
        d = docx.Document(bio)
        text = "\n".join(p.text for p in d.paragraphs)
        return title, _ensure_limit(_normalize_newlines(text), cfg.text_max_chars)

    if suffix in {".xlsx", ".xls"}:
        try:
            import pandas as pd
        except Exception as e:
            raise RuntimeError("Please install 'pandas' to parse Excel") from e
        xls = pd.ExcelFile(io.BytesIO(raw))
        out: List[str] = []
        for name in xls.sheet_names[:10]:
            df = pd.read_excel(xls, sheet_name=name, nrows=cfg.excel_preview_rows, engine=None)
            if df.shape[1] > cfg.excel_preview_cols:
                df = df.iloc[:, : cfg.excel_preview_cols]
            buf = io.StringIO()
            df.to_csv(buf, index=False, sep="\t")
            out.append(f"=== Sheet: {name} ===\n{buf.getvalue()}")
        return title, _ensure_limit(_normalize_newlines("\n\n".join(out)), cfg.text_max_chars)

    if suffix in {".html", ".htm"}:
        html = raw.decode("utf-8", errors="ignore")
        t, content = _load_html_from_string(html, prefer=cfg.html_extractor)
        return (t or title), _ensure_limit(_normalize_newlines(content), cfg.text_max_chars)

    raise ValueError(f"Unsupported file type: {suffix}")


async def load_text_from_url(url: str, *, config: Optional[LoaderConfig] = None) -> Tuple[str, str]:
    cfg = config or DEFAULT_CONFIG
    # 优先 trafilatura（自动抓取+正文抽取），失败则回退 httpx + 简单清洗
    try:
        import trafilatura
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            raise RuntimeError("trafilatura.fetch_url failed")

        title: Optional[str] = None
        try:
            from lxml import html as lxml_html
            doc = lxml_html.fromstring(downloaded)
            tnodes = doc.xpath("//title")
            if tnodes and tnodes[0].text:
                title = (tnodes[0].text or "").strip()
        except Exception:
            pass

        content = trafilatura.extract(
            downloaded,
            include_comments=False,
            include_tables=True,
            favor_recall=True,
        ) or ""
        if not content:
            # 再退：简单去标签
            content = _strip_html_basic(downloaded)

        title_str: str = title if (title and isinstance(title, str)) else url
        return title_str, _ensure_limit(_normalize_newlines(content), cfg.text_max_chars)

    except Exception:
        # 退化为 httpx 拉取 + 解析
        try:
            import httpx
            async with httpx.AsyncClient(timeout=20) as client:
                resp = await client.get(url)
                resp.raise_for_status()
                html = resp.text
        except Exception as e:
            raise RuntimeError(f"Failed to fetch url: {e}") from e

        # 你项目里已有该函数：返回 (title_optional, content)
        t, content = _load_html_from_string(html, prefer=cfg.html_extractor)
        title_str: str = t if (t and isinstance(t, str)) else url
        return title_str, _ensure_limit(_normalize_newlines(content), cfg.text_max_chars)
